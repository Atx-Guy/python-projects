from flask import Flask, request, render_template, send_file
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import markdown
import docx
from io import BytesIO
import re
from typing import Union
from bs4 import BeautifulSoup
import pypandoc
from rtf_parser import RTFParser

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

@app.route('/', methods=['GET'])
def index() -> str:
    """Renders the index page."""
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert() -> Union[str, send_file]:
    """Handles file conversion requests."""
    file = request.files.get('file')
    if not file:
        return "No file uploaded."

    filename = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(filename)

    input_format = request.form['input_format']
    output_format = request.form['output_format']
    return convert_file(filename, input_format, output_format)

def convert_file(
    filename: str, input_format: str, output_format: str
) -> Union[str, send_file]:
    """
    Converts a file from one format to another.

    Args:
        filename: The path to the input file.
        input_format: The format of the input file (e.g., '.txt', '.md').
        output_format: The desired output format (e.g., '.pdf', '.html').

    Returns:
        A file download response or an error message.
    """
    if input_format in ['.txt', '.html', '.md', '.rtf'] and output_format == '.pdf':
        return convert_text_to_pdf(filename)
    elif input_format == '.txt' and output_format == '.html':
        return convert_txt_to_html(filename)
    elif input_format == '.html' and output_format == '.txt':
        return convert_html_to_txt(filename)
    elif input_format == '.md' and output_format == '.html':
        return convert_md_to_html(filename)
    elif input_format in ['.txt', '.md'] and output_format == '.docx':
        return convert_text_to_docx(filename)
    elif input_format == '.docx' and output_format == '.txt':
        return convert_docx_to_txt(filename)
    elif input_format == '.txt' and output_format == '.rtf':
        return convert_txt_to_rtf(filename)
    elif input_format == '.rtf' and output_format == '.txt':
        return convert_rtf_to_txt(filename)
    elif input_format == '.rtf' and output_format == '.docx':
        return convert_rtf_to_docx(filename)
    elif input_format == '.txt' and output_format == '.md':
        return convert_txt_to_md(filename)
    elif input_format == '.md' and output_format == '.txt':
        return convert_md_to_txt(filename)
    elif input_format == '.docx' and output_format == '.pdf':
        return convert_docx_to_pdf(filename)
    elif input_format == '.docx' and output_format == '.rtf':
        return convert_docx_to_rtf(filename)
    elif input_format == '.docx' and output_format == '.md':
        return convert_docx_to_md(filename)
    elif input_format == '.html' and output_format == '.md':
        return convert_html_to_md(filename)
    elif input_format == '.html' and output_format == '.rtf':
        return convert_html_to_rtf(filename)
    elif input_format == '.html' and output_format == '.docx':
        return convert_html_to_docx(filename)
    elif input_format == '.md' and output_format == '.rtf':
        return convert_md_to_rtf(filename)
    #Add audio conversions here when pydub can be used.
    else:
        return "Conversion not supported yet."

def convert_txt_to_html(filename: str) -> Union[str, send_file]:
    """Converts a text file to HTML."""
    try:
        with open(filename, 'r') as f:
            content = f.read()
        html_content = f"<html><body><pre>{content}</pre></body></html>"
        output_filename = filename.replace('.txt', '.html')
        with open(output_filename, 'w') as f:
            f.write(html_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_html_to_txt(filename: str) -> Union[str, send_file]:
    """Converts an HTML file to plain text, keeping HTML tags intact."""
    try:
        with open(filename, 'r') as f:
            html_content = f.read()
        output_filename = filename.replace('.html', '.txt')
        with open(output_filename, 'w') as f:
            f.write(html_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_text_to_pdf(filename: str) -> Union[str, send_file]:
    """Converts a text file to PDF."""
    try:
        with open(filename, 'r') as f:
            content = f.read()
        output_filename = filename.rsplit('.', 1)[0] + '.pdf'
        c = canvas.Canvas(output_filename, letter)
        c.drawString(100, 750, content)
        c.save()
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_md_to_html(filename: str) -> Union[str, send_file]:
    """Converts a Markdown file to HTML."""
    try:
        with open(filename, 'r') as f:
            content = f.read()
        html_content = markdown.markdown(content)
        output_filename = filename.replace('.md', '.html')
        with open(output_filename, 'w') as f:
            f.write(html_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_text_to_docx(filename: str) -> Union[str, send_file]:
    """Converts a text or Markdown file to DOCX."""
    try:
        doc = docx.Document()
        with open(filename, 'r') as f:
            content = f.read()
        doc.add_paragraph(content)
        output_filename = filename.rsplit('.', 1)[0] + '.docx'
        doc.save(output_filename)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_docx_to_txt(filename: str) -> Union[str, send_file]:
    """Converts a DOCX file to plain text."""
    try:
        doc = docx.Document(filename)
        content = '\n'.join([p.text for p in doc.paragraphs])
        output_filename = filename.replace('.docx', '.txt')
        with open(output_filename, 'w') as f:
            f.write(content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_txt_to_rtf(filename: str) -> Union[str, send_file]:
    """Converts a text file to RTF."""
    try:
        with open(filename, 'r') as f:
            content = f.read()
        rtf_content = r"{\rtf1\ansi " + content.replace("\n", r"\line ") + r"}"
        output_filename = filename.replace('.txt', '.rtf')
        with open(output_filename, 'w') as f:
            f.write(rtf_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_rtf_to_txt(filename: str) -> Union[str, send_file]:
    """Converts an RTF file to plain text."""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            rtf_content = f.read()
        text_content = RTFParser.strip_rtf(rtf_content)
        output_filename = filename.replace('.rtf', '.txt')
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write(text_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_rtf_to_docx(filename: str) -> Union[str, send_file]:
    """Converts an RTF file to DOCX by first converting to plain text."""
    try:
        # First convert RTF to plain text
        with open(filename, 'r', encoding='utf-8') as f:
            rtf_content = f.read()
        text_content = RTFParser.strip_rtf(rtf_content)
        
        # Then use the existing text to DOCX conversion
        doc = docx.Document()
        doc.add_paragraph(text_content)
        output_filename = filename.replace('.rtf', '.docx')
        doc.save(output_filename)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_rtf_to_pdf(filename: str) -> Union[str, send_file]:
    """Converts an RTF file to PDF."""
    try:
        output_filename = filename.replace('.rtf', '.pdf')
        pypandoc.convert_file(filename, 'pdf', outputfile=output_filename)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_rtf_to_html(filename: str) -> Union[str, send_file]:
    """Converts an RTF file to HTML."""
    try:
        output_filename = filename.replace('.rtf', '.html')
        pypandoc.convert_file(filename, 'html', outputfile=output_filename)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_rtf_to_md(filename: str) -> Union[str, send_file]:
    """Converts an RTF file to Markdown."""
    try:
        output_filename = filename.replace('.rtf', '.md')
        pypandoc.convert_file(filename, 'markdown', outputfile=output_filename)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_md_to_txt(filename: str) -> Union[str, send_file]:
    """Converts a Markdown file to plain text."""
    try:
        with open(filename, 'r') as f:
            content = f.read()
        # Very basic conversion from Markdown to plain text.
        # Remove Markdown formatting patterns.
        text_content = re.sub(r'(!?\[.*?\]\(.*?\))', '', content)  # Remove images and links
        text_content = re.sub(r'(```.*?```)', '', text_content, flags=re.DOTALL)  # Remove code blocks
        text_content = re.sub(r'[`*_>#-]', '', text_content)  # Remove common markdown symbols
        output_filename = filename.replace('.md', '.txt')
        with open(output_filename, 'w') as f:
            f.write(text_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_txt_to_md(filename: str) -> Union[str, send_file]:
    """Converts a text file to Markdown (basic conversion)."""
    try:
        with open(filename, 'r') as f:
            content = f.read()
        # In this basic conversion, we assume the text is already Markdown friendly.
        output_filename = filename.rsplit('.', 1)[0] + '.md'
        with open(output_filename, 'w') as f:
            f.write(content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_docx_to_pdf(filename: str) -> Union[str, send_file]:
    """
    Converts a DOCX file to PDF.

    Args:
        filename: The path to the input DOCX file.

    Returns:
        A file download response or an error message.
    """
    try:
        doc = docx.Document(filename)
        output_filename = filename.replace('.docx', '.pdf')

        # Create a PDF canvas
        c = canvas.Canvas(output_filename, pagesize=letter)

        # Set initial y-coordinate for drawing text
        y = 750

        # Iterate through paragraphs and write content to PDF
        for paragraph in doc.paragraphs:
            # Add line breaks for newlines within paragraphs
            lines = paragraph.text.split('\n')
            for line in lines:
                c.drawString(100, y, line)
                y -= 15  # Adjust y-coordinate for next line

            # Add extra space between paragraphs
            y -= 10

        # Save the PDF
        c.save()

        # Send the PDF file as a download response
        return send_file(output_filename, as_attachment=True)

    except Exception as e:
        return f"Error: {e}"

def convert_docx_to_rtf(filename: str) -> Union[str, send_file]:
    """
    Converts a DOCX file to RTF.

    Args:
        filename: The path to the input DOCX file.

    Returns:
        A file download response or an error message.
    """
    try:
        doc = docx.Document(filename)
        output_filename = filename.replace('.docx', '.rtf')

        # Create a new RTF file and write the DOCX content to it
        with open(output_filename, 'w') as f:
            for paragraph in doc.paragraphs:
                f.write(paragraph.text + '\n')

        # Send the RTF file as a download response
        return send_file(output_filename, as_attachment=True)

    except Exception as e:
        return f"Error: {e}"

def convert_docx_to_md(filename: str) -> Union[str, send_file]:
    """
    Converts a DOCX file to Markdown.

    Args:
        filename: The path to the input DOCX file.

    Returns:
        A file download response or an error message.
    """
    try:
        doc = docx.Document(filename)
        output_filename = filename.replace('.docx', '.md')

        with open(output_filename, 'w') as f:
            for paragraph in doc.paragraphs:
                f.write(paragraph.text + '\n')

        return send_file(output_filename, as_attachment=True)

    except Exception as e:
        return f"Error: {e}"

def convert_html_to_md(filename: str) -> Union[str, send_file]:
    """Converts an HTML file to Markdown, keeping HTML tags intact."""
    try:
        with open(filename, 'r') as f:
            html_content = f.read()
        output_filename = filename.replace('.html', '.md')
        with open(output_filename, 'w') as f:
            f.write(html_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_html_to_rtf(filename: str) -> Union[str, send_file]:
    """Converts an HTML file to RTF, keeping HTML tags intact."""
    try:
        with open(filename, 'r') as f:
            html_content = f.read()
        rtf_content = r"{\rtf1\ansi " + html_content.replace("\n", r"\line ") + r"}"
        output_filename = filename.replace('.html', '.rtf')
        with open(output_filename, 'w') as f:
            f.write(rtf_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_html_to_docx(filename: str) -> Union[str, send_file]:
    """Converts an HTML file to DOCX, keeping HTML tags intact."""
    try:
        with open(filename, 'r') as f:
            html_content = f.read()
        doc = docx.Document()
        doc.add_paragraph(html_content)
        output_filename = filename.replace('.html', '.docx')
        doc.save(output_filename)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"

def convert_md_to_rtf(filename: str) -> Union[str, send_file]:
    """Converts a Markdown file to RTF."""
    try:
        with open(filename, 'r') as f:
            content = f.read()
        # Very basic conversion from Markdown to RTF.
        rtf_content = r"{\rtf1\ansi " + re.sub(r'[`*_>#-]', '', content).replace("\n", r"\line ") + r"}"
        output_filename = filename.replace('.md', '.rtf')
        with open(output_filename, 'w') as f:
            f.write(rtf_content)
        return send_file(output_filename, as_attachment=True)
    except Exception as e:
        return f"Error: {e}"